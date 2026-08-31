"""
IARD Render Service — self-hosted document assembly for the Internal Audit
Report Drafter Agent.

Replaces the n8n Document Assembly workflow's fabricated "microsoftWord" node.
n8n calls POST /render-report with the current engagement state (front matter +
observations + appendix items); this service fills Beyon's actual Word template
and returns the rendered .docx as bytes.

Design constraints carried over from the BRD / HLD:
 - GR-02 template fidelity: structure, table shapes, and rating colours come
   from the fixed template, never restyled by this service.
 - Verbatim sections (objectives, scope, limitations, risk classification) are
   inserted exactly as received — this service does not paraphrase anything.
 - Each observation paragraph gets a stable bookmark named obs-<observation_id>
   so a later re-uploaded, auditor-edited Word file can be diffed back to the
   right row (feeds the not-yet-built Round-trip Reconciliation sub-workflow).
 - area_for_improvement observations get no rating colour and are excluded
   from the Audit Issues Dashboard counts and the List of Audit Issues (FR-26).
"""

import io
import os
import base64
import copy
import logging
from typing import Optional

from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table

logger = logging.getLogger("iard-render-service")

# Beyon's fixed template is Restricted-classified (BRD Section 10) and is never committed to git
# (see Dockerfile). It's provided to the running container instead of being baked into the image.
#
# Render's Secret Files field is a text box, and a binary .docx pasted/uploaded through it directly
# does not survive intact — confirmed 2026-08-31: the Secret File existed at
# /etc/secrets/report_template.docx (os.path.isfile() true) but was empty, and even a non-empty
# direct paste risks the same silent corruption (line-ending normalization, encoding conversion,
# truncation). Fix: the template is base64-encoded first, so what actually goes in the text box is
# plain ASCII that survives a text field intact, and this service decodes it back to the real
# .docx byte-for-byte at startup.
#
# TEMPLATE_B64_SECRET_PATH: where the base64 Secret File is mounted (default matches Render's
#   /etc/secrets/<filename> convention for a file named report_template.b64).
# TEMPLATE_PATH: where the decoded, real .docx is written — must be writable (/etc/secrets itself
#   is read-only), hence /tmp by default.
#
# For local dev, skip the encoding: set TEMPLATE_PATH directly to a plain .docx file next to
# main.py and leave TEMPLATE_B64_SECRET_PATH pointing at a path that doesn't exist — decoding is
# skipped whenever TEMPLATE_PATH already resolves to a real, valid file.
#
# If a persistent Disk is used instead of a Secret File, skip the base64 dance entirely (Disks are
# binary-safe) — just set TEMPLATE_PATH to the Disk's mount path.
TEMPLATE_B64_SECRET_PATH = os.environ.get("TEMPLATE_B64_SECRET_PATH", "/etc/secrets/report_template.b64")
TEMPLATE_PATH = os.environ.get("TEMPLATE_PATH", "/tmp/report_template.docx")


def _template_is_valid(path: str) -> bool:
    if not os.path.isfile(path):
        return False
    try:
        Document(path)
        return True
    except Exception:
        return False


def ensure_template_decoded() -> None:
    """Decode the base64 Secret File into a real .docx at TEMPLATE_PATH, if needed.

    Idempotent and safe to call on every startup: does nothing if TEMPLATE_PATH already points at a
    valid docx (local dev with a plain .docx, or a Disk-mounted template — both binary-safe, no
    base64 involved; also covers a redeploy where the decoded file already exists from a prior
    boot). Does nothing if the base64 Secret File isn't present either (misconfigured deploy —
    /health's template_found/template_valid fields surface that instead of failing silently here).
    Re-decodes if TEMPLATE_PATH exists but isn't a valid docx (e.g. a leftover empty/corrupt file
    from before this fix), rather than treating "a file is there" as good enough.
    """
    if _template_is_valid(TEMPLATE_PATH):
        return
    if not os.path.isfile(TEMPLATE_B64_SECRET_PATH):
        logger.warning(
            "No valid template at TEMPLATE_PATH=%s and no base64 Secret File at "
            "TEMPLATE_B64_SECRET_PATH=%s — /render-report will fail until one is configured.",
            TEMPLATE_PATH, TEMPLATE_B64_SECRET_PATH,
        )
        return
    try:
        with open(TEMPLATE_B64_SECRET_PATH, "r") as f:
            b64_text = f.read().strip()
        docx_bytes = base64.b64decode(b64_text, validate=True)
        os.makedirs(os.path.dirname(TEMPLATE_PATH) or ".", exist_ok=True)
        with open(TEMPLATE_PATH, "wb") as f:
            f.write(docx_bytes)
        logger.info("Decoded template from %s to %s (%d bytes).", TEMPLATE_B64_SECRET_PATH, TEMPLATE_PATH, len(docx_bytes))
    except Exception as e:
        # Don't crash the whole service over a bad template — /health and /render-report both
        # surface this clearly instead, and everything else the service does still works.
        logger.error("Failed to decode base64 template from %s: %s", TEMPLATE_B64_SECRET_PATH, e)

RATING_LABELS = {
    "active_management_very_high": "Active Management (Very High)",
    "continuous_review_high": "Continuous Review (High)",
    "periodic_monitoring_medium": "Periodic Monitoring (Medium)",
    "no_major_concern_low": "No Major Concern (Low)",
    "area_for_improvement": "Area for Improvement",
}
# RGB fills matching the template's existing colour scheme for each rating.
RATING_COLOURS = {
    "active_management_very_high": "C00000",
    "continuous_review_high": "FF0000",
    "periodic_monitoring_medium": "FFC000",
    "no_major_concern_low": "92D050",
    "area_for_improvement": None,  # no colour, per FR-26
}
RISK_LETTER = {
    "active_management_very_high": "VH",
    "continuous_review_high": "H",
    "periodic_monitoring_medium": "M",
    "no_major_concern_low": "L",
    "area_for_improvement": "AFI",
}

app = FastAPI(title="IARD Render Service")


@app.on_event("startup")
def on_startup():
    ensure_template_decoded()


class RootCause(BaseModel):
    category_people: bool = False
    category_process: bool = False
    category_technology: bool = False
    explanation: Optional[str] = None
    flagged: bool = False  # evidence insufficient to support this specific cause (GR-04, now per-cause)


class Observation(BaseModel):
    id: str
    sequence_no: int
    title: Optional[str] = None
    criteria: Optional[str] = None
    condition: Optional[str] = None
    root_causes: list[RootCause] = []  # v0.2: plural, each categorised People/Process/Technology — replaces the old single `cause` string
    risk: Optional[str] = None
    recommendation: Optional[str] = None
    risk_rating: Optional[str] = None
    management_action: Optional[str] = None
    owner: Optional[str] = None
    target_date: Optional[str] = None


class AppendixItem(BaseModel):
    observation_id: str
    appendix_number: str
    content_ref: str


class FrontMatter(BaseModel):
    cover_page: dict
    objectives: Optional[list] = None
    scope: Optional[dict] = None
    limitations_of_internal_audit: Optional[str] = None
    risk_classification: Optional[str] = None
    acknowledgment: Optional[str] = None
    audit_team: Optional[dict] = None
    overall_opinion_paragraph: Optional[str] = None
    background: Optional[str] = None


class RenderRequest(BaseModel):
    engagement_id: str
    front_matter: FrontMatter
    observations: list[Observation]
    appendix_items: list[AppendixItem] = []


def add_bookmark(paragraph, bookmark_id: int, name: str):
    """Wrap a paragraph in a Word bookmark so a re-uploaded edit can be traced
    back to the observation it came from (round-trip reconciliation)."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def find_heading(doc: Document, text_startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith(text_startswith.lower()):
            return p
    return None


def insert_paragraph_after(paragraph, text: str = "", style=None, bold: bool = False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        if bold:
            run.bold = True
    if style:
        new_para.style = style
    return new_para


def root_cause_category_label(rc: "RootCause") -> str:
    cats = []
    if rc.category_people:
        cats.append("People")
    if rc.category_process:
        cats.append("Process")
    if rc.category_technology:
        cats.append("Technology")
    return ", ".join(cats)


def shade_cell(cell, hex_colour: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    tc_pr.append(shd)


def find_rating_card_tables(doc: Document) -> dict:
    """Locate the 5 per-rating "Detailed Audit Issues" card tables (one per RATING_LABELS value,
    e.g. "Active Management (Very High)"), by matching each table's row-0 last cell against the
    known rating labels — verified against the real template: row 0's first grid columns are a
    merged empty banner cell, and the last cell holds the rating text. Returns {rating_key: Table}.
    """
    label_to_key = {v: k for k, v in RATING_LABELS.items()}
    found = {}
    for table in doc.tables:
        if not table.rows:
            continue
        last_cell_text = table.rows[0].cells[-1].text.strip()
        rating_key = label_to_key.get(last_cell_text)
        if rating_key and rating_key not in found:
            found[rating_key] = table
    return found


def remove_table(table) -> None:
    """Delete a card table entirely — used when a rating has zero observations, so the issued
    report doesn't show an empty rating section that was only ever a template placeholder."""
    table._tbl.getparent().remove(table._tbl)


def set_cell_paragraphs(cell, lines: list) -> "Paragraph":
    """Write a list of (text, bold) lines into a table cell as separate paragraphs, reusing the
    cell's existing (usually empty) first paragraph for the first line rather than leaving a blank
    line above the content. Returns the first paragraph, for bookmarking."""
    first_para = cell.paragraphs[0]
    first_para.text = ""
    first_text, first_bold = lines[0] if lines else ("", False)
    run = first_para.add_run(first_text)
    run.bold = first_bold
    for text, bold in lines[1:]:
        p = cell.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
    return first_para


def render_report(req: RenderRequest) -> bytes:
    doc = Document(TEMPLATE_PATH)
    fm = req.front_matter

    # --- Cover page (verbatim + intake) ---
    cover = fm.cover_page
    for p in doc.paragraphs[:15]:
        t = p.text
        if "Project title" in t or "{{title}}" in t:
            p.text = cover.get("title", "")
        elif "Quarter and Year" in t or "{{quarter_year}}" in t:
            p.text = f"{cover.get('quarter', '')} {cover.get('year', '')}"
        elif "Draft Report" in t or "{{status}}" in t:
            p.text = cover.get("status", "Draft Report")

    # --- Background (generated) ---
    bg_heading = find_heading(doc, "Background")
    if bg_heading and fm.background:
        insert_paragraph_after(bg_heading, fm.background)

    # --- Objectives (verbatim) ---
    obj_heading = find_heading(doc, "Objectives:")
    if obj_heading and fm.objectives:
        anchor = obj_heading
        for item in fm.objectives:
            anchor = insert_paragraph_after(anchor, f"• {item}")

    # --- Scope (verbatim) ---
    scope_heading = find_heading(doc, "Scope:")
    if scope_heading and fm.scope:
        anchor = scope_heading
        for item in (fm.scope.get("items") or []):
            anchor = insert_paragraph_after(anchor, f"• {item}")
        if fm.scope.get("limitations"):
            anchor = insert_paragraph_after(anchor, fm.scope["limitations"])

    # --- Acknowledgment (fill) ---
    ack_heading = find_heading(doc, "Acknowledgment")
    if ack_heading and fm.acknowledgment:
        insert_paragraph_after(ack_heading, fm.acknowledgment)

    # --- Overall opinion (verbatim + selection) ---
    opinion_heading = find_heading(doc, "Overall opinion")
    if opinion_heading and fm.overall_opinion_paragraph:
        insert_paragraph_after(opinion_heading, fm.overall_opinion_paragraph)

    # --- Audit team (verbatim + fill) ---
    if fm.audit_team:
        for table in doc.tables:
            header_text = " ".join(c.text for c in table.rows[0].cells).lower()
            if "audit team" in header_text or len(table.rows) <= 3 and table.columns and len(table.columns) == 3:
                # best-effort: fill first data row with auditor/audit_manager/chief names
                try:
                    row = table.rows[1] if len(table.rows) > 1 else table.rows[0]
                    row.cells[0].text = fm.audit_team.get("auditor", "")
                    row.cells[1].text = fm.audit_team.get("audit_manager", "")
                    row.cells[2].text = fm.audit_team.get("chief_of_internal_audit", "")
                except IndexError:
                    pass
                break

    # --- Audit Issues Dashboard + List of Audit Issues (derived; exclude area_for_improvement) ---
    listable = [o for o in req.observations if o.risk_rating != "area_for_improvement"]

    dashboard_table = None
    list_table = None
    for table in doc.tables:
        header_text = " ".join(c.text for c in table.rows[0].cells).lower()
        if "business area" in header_text:
            dashboard_table = table
        if header_text.strip().startswith("n") and "audit issue" in header_text:
            list_table = table

    if list_table:
        # The template ships with sample rows already filled in under the header (placeholder
        # "New Audit Issue" text with an H/H/M/M/M/L pattern) — verified against the real template
        # 2026-08-31. Those aren't blank placeholders to leave for later like the dashboard's "-"
        # cells; they're leftover sample content and must be cleared before real rows go in, or
        # every render would mix real observations with fake ones.
        for row in list(list_table.rows[1:]):
            list_table._tbl.remove(row._tr)
        for idx, obs in enumerate(listable, start=1):
            row_cells = list_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = RISK_LETTER.get(obs.risk_rating, "")
            row_cells[2].text = obs.title or ""
            # Owner/target date deliberately left blank in draft (FR-16 / content rule)

    if dashboard_table:
        # Only the Total row is genuinely computable from what this service receives — the
        # Business Area breakdown rows require a per-observation business-area field that doesn't
        # exist in the Observation model or the DB schema (001/002 schema review confirms this), so
        # inventing area names or splitting counts across them would violate GR-04 (don't invent
        # what evidence doesn't support). Those rows are left exactly as the template has them
        # (their own "-" placeholder, for the auditor to fill in manually), and only the Total
        # row's dashes are replaced with real counts by rating, columns matching the template's own
        # header order: Very High / High / Medium / Low / Total.
        rating_to_col = {
            "active_management_very_high": 1,
            "continuous_review_high": 2,
            "periodic_monitoring_medium": 3,
            "no_major_concern_low": 4,
        }
        counts = {col: 0 for col in rating_to_col.values()}
        for obs in listable:
            col = rating_to_col.get(obs.risk_rating)
            if col is not None:
                counts[col] += 1
        total_row = dashboard_table.rows[-1]
        for col in (1, 2, 3, 4):
            total_row.cells[col].text = str(counts[col])
        total_row.cells[5].text = str(sum(counts.values()))

    # --- Detailed Audit Issues (generated; one card per observation) ---
    # The template ships with exactly one blank "card" table per rating (verified against the real
    # template 2026-08-31 — NOT a heading-plus-paragraphs section as previously assumed). Each card's
    # last row holds two content cells: "Audit Issue" and "Management Action" (a third, narrow column
    # is a decorative colour strip, merged into the Management Action cell on the 4 rated cards, and
    # confirmed to carry no data on any card — including Area for Improvement, which has no separate
    # "Audit Issue" label row at all). A rating with zero observations has its card removed entirely;
    # a rating with more than one observation gets its card cloned once per extra observation, stacked
    # in sequence_no order.
    card_tables = find_rating_card_tables(doc)
    bookmark_counter = 1

    observations_by_rating = {}
    for obs in sorted(req.observations, key=lambda o: o.sequence_no):
        observations_by_rating.setdefault(obs.risk_rating, []).append(obs)

    for rating_key, template_table in card_tables.items():
        obs_list = observations_by_rating.get(rating_key, [])
        if not obs_list:
            remove_table(template_table)
            continue

        # Capture the card's blank XML before any filling happens — clones must come from this
        # pristine copy, not from a sibling card that's already been filled (which would carry its
        # content into the "blank" clone).
        blank_tbl_xml = copy.deepcopy(template_table._tbl)
        prev_card = template_table
        for i, obs in enumerate(obs_list):
            if i == 0:
                card = template_table
            else:
                new_tbl = copy.deepcopy(blank_tbl_xml)
                prev_card._tbl.addnext(new_tbl)
                card = Table(new_tbl, prev_card._parent)
            prev_card = card
            content_row = card.rows[-1]
            issue_cell, action_cell = content_row.cells[0], content_row.cells[1]

            lines = [(f"{obs.sequence_no}. {obs.title or '[untitled]'}", True)]
            if obs.criteria:
                lines.append((f"Criteria: {obs.criteria}", False))
            if obs.condition:
                lines.append((f"Condition: {obs.condition}", False))
            # v0.2: root causes are plural, each categorised People/Process/Technology (or a
            # combination), each with its own explanation. Verified against the real template and
            # issued CVM report — there is no precedent there for multiple/tagged causes (that's new
            # in v0.2), so this keeps the template's existing convention (one bold "Potential Cause"
            # label, plain text under it) rather than inventing a new heading or table, and lists one
            # cause per line under that label. Confirmed with Nuella 2026-08-25.
            if obs.root_causes:
                lines.append(("Potential Cause", True))
                for rc in obs.root_causes:
                    if rc.flagged or not rc.explanation:
                        lines.append(("[FLAGGED — evidence insufficient to support this cause; auditor input required]", False))
                    else:
                        rc_label = root_cause_category_label(rc)
                        prefix = f"({rc_label}) — " if rc_label else ""
                        lines.append((f"{prefix}{rc.explanation}", False))
            elif obs.risk_rating != "area_for_improvement":
                lines.append(("Potential Cause", True))
                lines.append(("[FLAGGED — evidence insufficient to support a cause; auditor input required]", False))
            if obs.risk:
                lines.append((f"Risk: {obs.risk}", False))
            if obs.recommendation:
                lines.append((f"Recommendation: {obs.recommendation}", False))

            appx = [a for a in req.appendix_items if a.observation_id == obs.id]
            if appx:
                refs = "; ".join(f"Appendix {a.appendix_number}: {a.content_ref}" for a in appx)
                lines.append((f"See: {refs}", False))

            first_para = set_cell_paragraphs(issue_cell, lines)
            add_bookmark(first_para, bookmark_counter, f"obs-{obs.id}")
            bookmark_counter += 1

            # Management Action: never written by the agent — only rendered if the auditor already entered it.
            action_cell.text = obs.management_action or ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.post("/render-report")
def render_report_endpoint(req: RenderRequest):
    try:
        docx_bytes = render_report(req)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"render failed: {e}")
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{req.engagement_id}.docx"'},
    )


@app.get("/health")
def health():
    # Surfaces the template-availability problem directly, instead of only discovering it on the
    # first real /render-report call (which is how we found it — Aug 28/31 failures both traced
    # back to this). Still 200s so this doesn't flap the service's deploy health check; the
    # template_valid field is what to check.
    #
    # os.path.isfile() alone isn't enough here — it only proves *something* exists at TEMPLATE_PATH,
    # not that it's a real, openable .docx. python-docx raises the same "Package not found at
    # '<path>'" error both when the file is missing AND when it exists but isn't a valid Word/zip
    # package — e.g. corrupted by however it got uploaded (Render's Secret File field is a text
    # paste box; a binary .docx pasted/uploaded through it can come out mangled even though a file
    # does land at the path). A run on 2026-08-31 hit exactly this: /health reported the file
    # present right before a /render-report call failed with that same error. So this actually
    # tries to open it, the same way render_report() will.
    exists = os.path.isfile(TEMPLATE_PATH)
    valid = False
    error = None
    if exists:
        try:
            Document(TEMPLATE_PATH)
            valid = True
        except Exception as e:
            error = f"{type(e).__name__}: {e}"
    return {
        "status": "ok",
        "template_path": TEMPLATE_PATH,
        "template_found": exists,
        "template_valid": valid,
        "template_error": error,
        "template_b64_secret_path": TEMPLATE_B64_SECRET_PATH,
        "template_b64_secret_found": os.path.isfile(TEMPLATE_B64_SECRET_PATH),
    }
